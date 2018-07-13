
# coding: utf-8

# ## Text Extraction
#
# ##### Author: Daniel Beaulieu | danbeauleu@gmail.com
#
#
#Purpose of program
# - Extract Text from Word Documents
# - Identify style (e.g. Bold, Font) and metadata (e.g. author) associated with document text
# - Understand docx XML tag definitions
# - Learn how to interact with Zip Files
# - Identify content surrounding key piece of text
# - Extract text from a pdf with pdfminer.six

import os
from IPython.core.display import display, HTML
from configparser import ConfigParser, ExtendedInterpolation

config = ConfigParser(interpolation=ExtendedInterpolation())
config.read('../../config.ini')

DOCX_PATH = config['DOCX']['DOCX_PATH']
XML_PATH = config['DOCX']['XML_PATH']
EXAMPLE_ZIP = config['DOCX']['EXAMPLE_ZIP']


# ### python-docx
#
# python-docx is a Python library for  extracting text from Microsoft Word (.docx) files.
# the Document method reads the text, style, and formatting of a word .docx document
import docx
doc = docx.Document(DOCX_PATH)
# view the methods and attributes of a doc
print(dir(doc))


# ### Paragraphs
paragraphs = doc.paragraphs
# view the docx paragraph objects
paragraphs[0:5]
# count all paragraphs in the document
len(paragraphs)

# only include with text (ignore empty strings)
paragraphs = [p for p in paragraphs if p.text.strip() != '']

# view the text of the first paragraph
paragraphs[0].text

# ### Style
# view the methods and attributes of a paragraph
print(dir(paragraphs[0]))

# get the paragraph style
paragraphs[0].style.name

# Identify if paragraph text has 'Heading' style
'heading' in paragraphs[0].style.name.lower()

# view all the heading styles in the doc
set(p.style.name for p in paragraphs if 'heading' in p.style.name.lower())

# store all heading paragraphs
headings = [p.text.strip() for p in paragraphs if 'heading' in p.style.name.lower()]

print('# heading paragraphs: {}\n'.format(len(headings)))
headings[0:10]


# ### Runs
# Each paragraph may contain one or more runs. A run denotes the style attached to the text in a paragraph. Every time the style change (e.g. from bold to normal text) a new run is added.

runs = paragraphs[0].runs
runs

# View all the runs in the paragraph
[run.text for run in runs]
# each run contains a portion of text from the paragraph
run = runs[2]
run.text

# ### Run style
#
# - Each run contains style information such as bold, italic, or underline.
# - The style information will be True, False, or None

# view the methods and attributes of a run
print(dir(run))

# font size
run.font.size.pt
print(run.bold)
print(run.italic)
print(run.underline)

# View all the run stlye
[run.bold for run in runs]

# #### Find all the bold runs

bold_text = []
for paragraph in paragraphs:
    for run in paragraph.runs:
        if run.bold and run.text.strip() != '':
            text = run.text
            bold_text.append(text)

bold_text[0:10]


# ### Create a function to determine if all runs in a paragraph are bold
#
# create the function is_bold
def is_bold(paragraph):
    runs_are_bold = [run.bold for run in paragraph.runs if run.text != '']

    if runs_are_bold and all(runs_are_bold):  # runs_are_bold evaluates as False if the list is empty
        return True
    return False

# test the is_bold function
bold_paragraphs = []
for paragraph in paragraphs:
    if is_bold(paragraph):
        bold_paragraphs.append(paragraph.text)

bold_paragraphs[0:10]


# ### Tables
# identify all document tables
tables = doc.tables

# view a few table objects
tables[0:5]

# count the document tables
len(tables)

# view the methods and attributes of a table
print(dir(tables[0]))

# view the cells of a table

table_cells = [cell.text.strip() for cell in tables[0]._cells if cell.text != '']
table_cells[0:10]


# ### Core Properties
print(dir(doc.core_properties))
doc.core_properties.title
doc.core_properties.subject
doc.core_properties.author
doc.core_properties.created
doc.core_properties.revision

# ## Explore docx xml
# Every word document is a zip of xml files. To test this, change the extension of any word file from .docx to .xml.
# Inside each zip, a directory named word contains document.xml. This file contains all of the xml for the word document.
XML_PATH


# ### zipfile
#
# ZipFile - The class for reading and writing ZIP files
# read - Returns the bytes content from a zipfile
import zipfile

zipf = zipfile.ZipFile(XML_PATH, 'r')

for f in zipf.filelist:
    print(f.filename)

xml_content = zipf.read('word/document.xml')

from bs4 import BeautifulSoup

b = BeautifulSoup(xml_content, 'lxml')


# view the xml from a short document with one heading and one sentence
for word in b.find('w:body'):
    print(word)
    print()


# ### docx XML tag definitions
# - < w:body > - contains the document paragraphs
# - < w:p > - Document paragraph
# - < w:pstyle > Document Style (e.g. Header 1)
# - < w:t > text in a paragraph or run
# - < w:bookmarkstart > defines a bookmark, such as a link in a table of contents
# - < w:r > - Document runs. Every time the style in a paragraph changes, for instance a bold or underline term, a new run is added. Each paragraph may contain multiple runs.
#
# view the  directory - notice there is no 'word' directory
get_ipython().run_line_magic('ls', '')

# Extract a member from the archive to the current working directory
# view the directory with a 'word' directory
get_ipython().run_line_magic('ls', '')


#search through several Oracle annual reports to find selected text throughout all the documents without needing to extract the files from the zip manually.
XAMPLE_ZIP
 use zipfile to read the example_zip
zipf = zipfile.ZipFile(EXAMPLE_ZIP, 'r')
# How many documents are in the provided zip?
len(zipf.filelist)
# view the filenames
# use the .filename attribute on each file in zip.filelist
[f.filename for f in zipf.filelist]

# Find  paragraphs scattered in all the documents in the zip
# that speak about 'Financial Accounting Standards No. 109'

# iterate through the filelist
for f in zipf.filelist:
    # use zip.extract the file to the currect working directory
    doc_file = zipf.extract(f)
    # open the document with docx
    doc = docx.Document(doc_file)
    # iterate through the paragraphs in the document
    for p in doc.paragraphs:
        # check which paragraphs contain 'Financial Accounting Standards No. 109'
        if 'Financial Accounting Standards No. 109' in p.text:
            # print the paragraphs that meet the condition
            print(p.text)
            print()


# # PDF Text Extraction

# ##### subprocess - use python to interact with the command line
# run an ls from python
import subprocess
output = subprocess.check_output('dir', shell=True)
output.split()


# view  pdf in raw_data dir to extract text from using pdfminer.six
output = subprocess.check_output(['dir','raw_data'], shell=True)
output.split()


# ### pdfminer.six
#
# ##### Installation
# - conda install -c conda-forge pdfminer.six
#
# The pdf2txt.py command:
# - The package includes the pdf2txt.py command-line command, which you can use to extract text and images. The command supports many options and is very flexible. Some popular options are shown below. See the usage information for complete details.
#
# **pdf2txt.py [options] filename.pdf**
#
# Options:
# - o output file name
# - p comma-separated list of page numbers to extract
# - t output format (text/html/xml/tag[for Tagged PDFs])
# - O dirname (triggers extraction of images from PDF into directory)
# - P password
# add your username to read the local pdf
username = 'dabeaulieu'
# extract the first three pages of the pdf, output to a .txt
cmd = [
    'python'
  , r'C:\Users\{}\AppData\Local\Continuum\anaconda3\Scripts\pdf2txt.py'.format(username)  # pdfminer
  , 'raw_data\southwest-airlines-co_annual_report_2016.pdf '  # imput pdf
  , '-o'  # output file name
  , 'raw_data\southwest_2016.txt'
  , '-t'  # output format
  , 'text'
  , '-p'  # pages to extract, default is to extract all pages
  , '1,2,3'
]

subprocess.call(cmd, shell=True)


# check the raw_data dir for the extracted text from the pdf
output = subprocess.check_output(['dir','raw_data'], shell=True)
output.split()
